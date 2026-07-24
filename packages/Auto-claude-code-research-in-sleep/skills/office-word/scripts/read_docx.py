#!/usr/bin/env python3
# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0
"""Read a Word (.docx) file and output its content as Markdown."""
import sys
import os


def read_docx_to_markdown(docx_path: str) -> str:
    """Read a .docx file and convert to markdown with inline tables."""
    from docx import Document
    from docx.oxml.ns import qn

    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)
    lines = []

    # Iterate over body elements in document order to keep tables inline
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("Heading 4"):
                lines.append(f"#### {text}")
            elif style.startswith("List Bullet") or style.startswith("List Number"):
                lines.append(f"- {text}")
            else:
                parts = []
                for run in para.runs:
                    t = run.text
                    if not t:
                        continue
                    if run.bold and run.italic:
                        parts.append(f"***{t}***")
                    elif run.bold:
                        parts.append(f"**{t}**")
                    elif run.italic:
                        parts.append(f"*{t}*")
                    else:
                        parts.append(t)
                lines.append("".join(parts) if parts else text)

        elif tag == "tbl":
            tbl_obj = None
            for t in doc.tables:
                if t._element is element:
                    tbl_obj = t
                    break
            if tbl_obj is None:
                continue

            lines.append("")
            for row_idx, row in enumerate(tbl_obj.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: read_docx.py <input.docx> [output.md]", file=sys.stderr)
        sys.exit(1)

    docx_path = sys.argv[1]
    result = read_docx_to_markdown(docx_path)

    if len(sys.argv) > 2:
        with open(sys.argv[2], "w", encoding="utf-8") as f:
            f.write(result)
        print(f"Markdown saved to: {sys.argv[2]}")
    else:
        print(result)
