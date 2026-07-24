#!/usr/bin/env python3
# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0
"""Create a Word (.docx) file from Markdown content."""
import sys
import os
import re


def markdown_content_to_docx(content: str, output_path: str):
    """Convert markdown content string to .docx format."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    in_table = False
    table_rows = []

    for line in content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^-+$", c) for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols)
                table.style = "Table Grid"
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = cell_text
            table_rows = []
            in_table = False

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        elif stripped:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", stripped)
            for part in parts:
                if part.startswith("***") and part.endswith("***"):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    table.rows[i].cells[j].text = cell_text

    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


def markdown_to_docx(md_path: str, output_path: str):
    """Convert a markdown file to .docx format."""
    if not os.path.exists(md_path):
        print(f"Error: File not found: {md_path}", file=sys.stderr)
        sys.exit(1)
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
    markdown_content_to_docx(content, output_path)


if __name__ == "__main__":
    # Mode 1: stdin -> docx
    if len(sys.argv) >= 3 and sys.argv[1] == "--stdin":
        output_path = sys.argv[2]
        md_content = sys.stdin.read()
        if not md_content.strip():
            print("Error: No content received from stdin", file=sys.stderr)
            sys.exit(1)
        markdown_content_to_docx(md_content, output_path)
    # Mode 2: file -> docx
    elif len(sys.argv) >= 3:
        markdown_to_docx(sys.argv[1], sys.argv[2])
    else:
        print("Usage:", file=sys.stderr)
        print("  write_docx.py <input.md> <output.docx>", file=sys.stderr)
        print("  echo 'md content' | write_docx.py --stdin <output.docx>", file=sys.stderr)
        sys.exit(1)
